"""
Audit Assistant Agent - AG-UI endpoint using Pydantic AI.

This creates an AG-UI compatible endpoint that CopilotKit can connect to
for bidirectional state synchronization and streaming agent responses.
Also provides a file upload endpoint for document processing.

uv run uvicorn main:app --reload --port 8001
"""

import os
import re
import json
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Any
from uuid import uuid4
from dotenv import load_dotenv

# Load environment variables first
load_dotenv()

from fastapi import FastAPI, File, UploadFile, Body  # noqa: E402
from fastapi.middleware.cors import CORSMiddleware  # noqa: E402
from fastapi.staticfiles import StaticFiles  # noqa: E402
from pydantic import BaseModel, ConfigDict  # noqa: E402
from starlette.requests import Request  # noqa: E402
from starlette.responses import JSONResponse, StreamingResponse  # noqa: E402

from agent import agent, AuditState  # noqa: E402
from pydantic_ai.ui import StateDeps  # noqa: E402
from pydantic_ai.ui.ag_ui import AGUIAdapter  # noqa: E402
from a2ui_generator import generate_audit_question_form  # noqa: E402
from llm_orchestrator import (  # noqa: E402
    document_summary_agent,
    batch_tagger_agent,
    search_sort_agent,
    SearchSortDeps,
)
from models import ALL_DOC_TAGS, Document, Documents  # noqa: E402
from prompts import (  # noqa: E402
    format_document_summary_prompt,
    format_batch_tagger_prompt,
    format_doc_search_sort_prompt,
)

# Configuration
BACKEND_PORT = int(os.getenv("BACKEND_PORT", "8001"))
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
FORMS_DIR = os.path.join(os.path.dirname(__file__), "data", "forms")
os.makedirs(FORMS_DIR, exist_ok=True)


# Shared in-memory state for AG-UI and persistence endpoints.
# This keeps UI edits and backend submit/restore operations in sync.
APP_DEPS = StateDeps(AuditState())


# =========================================================================
# Request body models
# =========================================================================


class AuditFormRequestBody(BaseModel):
    """Flexible audit form request body for state sync and persistence endpoints.

    Accepts the form payload in three formats:
    1. Wrapped in an ``audit_form_result`` envelope.
    2. Wrapped in a ``form`` envelope.
    3. Flat/direct payload fields passed at the top level (captured via ``extra='allow'``).

    Example usage::

        # Wrapped format
        body = AuditFormRequestBody(audit_form_result={"peril": {...}, "questions": [...]})

        # Direct format (extra fields)
        body = AuditFormRequestBody(**{"peril": {...}, "questions": [...], "overall_outcome": "Pass"})

        payload = body.extract_form_payload()
    """

    model_config = ConfigDict(extra="allow")

    audit_form_result: dict[str, Any] | None = None
    form: dict[str, Any] | None = None
    current_form_id: str | None = None
    id: str | None = None
    title: str | None = None
    source_docs: list[Any] = []

    def extract_form_payload(self) -> dict[str, Any]:
        """Extract canonical audit form payload from the request body.

        Checks for an ``audit_form_result`` or ``form`` wrapper first, then
        falls back to any extra fields passed directly at the top level.

        Returns:
            Dict containing the audit form data, or an empty dict if none found.
        """
        if isinstance(self.audit_form_result, dict):
            return self.audit_form_result
        if isinstance(self.form, dict):
            return self.form
        # Direct payload passed as extra top-level fields
        return self.model_extra or {}


# =========================================================================
# Helper utilities
# =========================================================================


def _iso_now() -> str:
    """Return the current UTC timestamp in ISO 8601 format."""
    return datetime.now(timezone.utc).isoformat()


def _atomic_write_json(path: str, payload: dict[str, Any]) -> None:
    """Write JSON atomically by writing to temp then replacing target file.

    Args:
        path: Destination JSON file path.
        payload: JSON-serializable dictionary payload.
    """
    tmp_path = f"{path}.tmp"
    with open(tmp_path, "w", encoding="utf-8") as file_obj:
        json.dump(payload, file_obj, indent=2, ensure_ascii=False)
    os.replace(tmp_path, path)


def _form_file_path(form_id: str) -> str:
    """Build a safe JSON file path for a form ID.

    Args:
        form_id: Form identifier.

    Returns:
        Full path to the form JSON file.
    """
    safe_id = re.sub(r"[^a-zA-Z0-9_-]", "_", form_id)
    return os.path.join(FORMS_DIR, f"{safe_id}.json")


def _validate_form_payload(payload: dict[str, Any]) -> str | None:
    """Validate required audit form fields for persistence and restore.

    Args:
        payload: Candidate audit form payload.

    Returns:
        None when valid, otherwise an error string.
    """
    required_fields = [
        "peril",
        "questions",
        "overall_outcome",
        "outcome_justification",
    ]
    missing = [field for field in required_fields if field not in payload]
    if missing:
        return f"Missing required fields: {', '.join(missing)}"

    if not isinstance(payload.get("questions"), list):
        return "Field 'questions' must be a list."

    if not isinstance(payload.get("peril"), dict):
        return "Field 'peril' must be an object."

    return None


def _build_form_title(form_payload: dict[str, Any]) -> str:
    """Build a human-friendly fallback title for a saved form.

    Args:
        form_payload: Canonical audit form payload.

    Returns:
        Generated title string.
    """
    peril = form_payload.get("peril", {}).get("peril", "Unknown")
    outcome = form_payload.get("overall_outcome", "Unknown")
    return f"{peril} - {outcome} - {_iso_now()[:10]}"


def _upsert_audit_form_component(state: AuditState, form_payload: dict[str, Any]) -> None:
    """Replace existing AuditQuestionForm component with latest payload.

    Args:
        state: Shared audit state.
        form_payload: Canonical audit form payload.
    """
    component = generate_audit_question_form(
        peril=form_payload["peril"],
        questions=form_payload["questions"],
        overall_outcome=form_payload["overall_outcome"],
        outcome_justification=form_payload["outcome_justification"],
        additional_analysis=form_payload.get("additional_analysis"),
        follow_ups=form_payload.get("follow_ups"),
    ).model_dump()

    # Keep only one active audit form component for predictable restore behavior.
    replaced = False
    for index, item in enumerate(state.components):
        if item.get("type") == "a2ui.AuditQuestionForm":
            state.components[index] = component
            replaced = True
            break
    if not replaced:
        state.components.append(component)


# =========================================================================
# App factory & lifespan
# =========================================================================


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application startup/shutdown lifecycle handler."""
    print(f"[*] Audit Assistant Agent (AG-UI) starting on port {BACKEND_PORT}")
    print(f"[*] AG-UI endpoint: POST http://localhost:{BACKEND_PORT}/")
    print(f"[*] Upload endpoint: POST http://localhost:{BACKEND_PORT}/upload")
    print(f"[*] Form state sync endpoint: GET/PUT http://localhost:{BACKEND_PORT}/state/audit-form")
    print(f"[*] Form persistence endpoints: POST/GET http://localhost:{BACKEND_PORT}/forms")
    print(
        f"[*] Form restore endpoint: POST http://localhost:{BACKEND_PORT}/forms/{{form_id}}/restore"
    )
    print(f"[*] Summarize endpoint: POST http://localhost:{BACKEND_PORT}/summarize (SSE)")
    print(f"[*] Info endpoint: GET http://localhost:{BACKEND_PORT}/info")
    print(f"[*] Health endpoint: GET http://localhost:{BACKEND_PORT}/health")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("[!] WARNING: OPENAI_API_KEY not set")
    else:
        print("[+] OpenAI API key configured")

    yield  # Application runs here; add shutdown logic after yield if needed


app = FastAPI(
    title="Audit Assistant Agent",
    description="Analyze documents and generate custom audit questionnaires",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================================================================
# AG-UI endpoint
# =========================================================================


@app.post("/")
async def ag_ui_endpoint(request: Request):
    """AG-UI endpoint backed by Pydantic AI's official request dispatcher."""
    try:
        body = await request.body()
        print(f"[AG-UI] Received request: {len(body)} bytes", flush=True)

        original_response = await AGUIAdapter.dispatch_request(
            request=request,
            agent=agent,
            deps=APP_DEPS,
        )

        return original_response
    except Exception as e:
        print(f"[AG-UI ERROR] {e}", flush=True)
        import traceback

        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/")
async def root_get():
    """Return AG-UI endpoint info for GET requests."""
    return JSONResponse(
        {
            "protocol": "ag-ui",
            "version": "1.0.0",
            "endpoints": {
                "run_agent": "POST /",
                "upload": "POST /upload",
                "summarize": "POST /summarize (SSE)",
                "audit_form_state": "GET|PUT /state/audit-form",
                "save_form": "POST /forms",
                "list_forms": "GET /forms",
                "get_form": "GET /forms/{form_id}",
                "restore_form": "POST /forms/{form_id}/restore",
                "info": "GET /info",
                "health": "GET /health",
            },
            "description": "Audit Assistant Agent - POST to / to run the agent",
        }
    )


# =========================================================================
# Audit form state sync + persistence endpoints
# =========================================================================


@app.get("/state/audit-form")
async def get_audit_form_state():
    """Get the current audit form payload and form ID from shared state."""
    state = APP_DEPS.state
    return JSONResponse(
        {
            "current_form_id": state.current_form_id,
            "audit_form_result": state.audit_form_result,
        }
    )


@app.put("/state/audit-form")
async def put_audit_form_state(body: AuditFormRequestBody):
    """Update the editable audit form payload in shared state.

    Upserts payload into state and keeps the AuditQuestionForm component in sync.
    Accepts the payload wrapped in ``audit_form_result``/``form``, or flat at the
    top level.
    """
    state = APP_DEPS.state

    payload = body.extract_form_payload()
    validation_error = _validate_form_payload(payload)
    if validation_error:
        return JSONResponse({"error": validation_error}, status_code=400)

    state.audit_form_result = payload
    state.audit_questions = payload["questions"]
    if body.current_form_id:
        state.current_form_id = body.current_form_id
    _upsert_audit_form_component(state, payload)

    return JSONResponse(
        {
            "message": "Audit form state synchronized.",
            "current_form_id": state.current_form_id,
            "audit_form_result": state.audit_form_result,
        }
    )


@app.post("/forms")
async def save_form(body: AuditFormRequestBody | None = Body(default=None)):
    """Persist an audit form to local JSON storage under data/forms/.

    Request body can include either:
    - audit_form_result/form payload directly
    - optional id/title/source_docs metadata

    If no valid payload is provided in the request body, this endpoint falls
    back to the current in-memory state audit_form_result.
    """
    state = APP_DEPS.state

    payload = body.extract_form_payload() if body else {}
    if not payload:
        # Fall back to whatever is currently in memory
        payload = state.audit_form_result

    validation_error = _validate_form_payload(payload)
    if validation_error:
        return JSONResponse({"error": validation_error}, status_code=400)

    requested_id = body.id if body else None
    form_id = requested_id or state.current_form_id or str(uuid4())
    path = _form_file_path(form_id)

    # Preserve created_at when updating an existing saved form.
    existing_created_at = None
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as file_obj:
                existing = json.load(file_obj)
                existing_created_at = existing.get("created_at")
        except Exception:
            existing_created_at = None

    record = {
        "id": form_id,
        "schema_version": "1.0",
        "created_at": existing_created_at or _iso_now(),
        "updated_at": _iso_now(),
        "title": (body.title if body and body.title else None) or _build_form_title(payload),
        "source_docs": body.source_docs if body else [],
        "peril": payload["peril"],
        "questions": payload["questions"],
        "overall_outcome": payload["overall_outcome"],
        "outcome_justification": payload["outcome_justification"],
        "additional_analysis": payload.get("additional_analysis"),
        "follow_ups": payload.get("follow_ups"),
    }

    _atomic_write_json(path, record)

    # Keep memory state aligned with saved payload.
    state.current_form_id = form_id
    state.audit_form_result = payload
    state.audit_questions = payload["questions"]
    _upsert_audit_form_component(state, payload)

    return JSONResponse(
        {
            "message": "Form saved.",
            "form_id": form_id,
            "title": record["title"],
            "updated_at": record["updated_at"],
        }
    )


@app.get("/forms")
async def list_forms():
    """List all saved forms from local JSON storage."""
    forms: list[dict[str, Any]] = []
    for name in os.listdir(FORMS_DIR):
        if not name.endswith(".json"):
            continue
        file_path = os.path.join(FORMS_DIR, name)
        try:
            with open(file_path, "r", encoding="utf-8") as file_obj:
                data = json.load(file_obj)
            forms.append(
                {
                    "id": data.get("id"),
                    "title": data.get("title"),
                    "created_at": data.get("created_at"),
                    "updated_at": data.get("updated_at"),
                    "peril": data.get("peril", {}).get("peril"),
                    "overall_outcome": data.get("overall_outcome"),
                    "question_count": len(data.get("questions", [])),
                }
            )
        except Exception as exc:
            # Skip unreadable/corrupt files but expose useful debug signal.
            print(f"[FORMS] Failed reading {file_path}: {exc}", flush=True)

    forms.sort(key=lambda form: form.get("updated_at") or "", reverse=True)
    return JSONResponse({"forms": forms})


@app.get("/forms/all")
async def list_forms_full():
    """Return all saved forms with full data for dashboard aggregation.

    Unlike ``GET /forms`` which returns lightweight summaries, this endpoint
    reads every JSON file in full so the frontend can compute cross-form
    aggregations (question-level stats, driver counts, etc.).

    Returns:
        JSON ``{"forms": [<full form record>, ...]}`` sorted newest-first.
    """
    forms: list[dict[str, Any]] = []
    for name in os.listdir(FORMS_DIR):
        if not name.endswith(".json"):
            continue
        file_path = os.path.join(FORMS_DIR, name)
        try:
            with open(file_path, "r", encoding="utf-8") as file_obj:
                data = json.load(file_obj)
            forms.append(data)
        except Exception as exc:
            print(f"[FORMS] Failed reading {file_path}: {exc}", flush=True)

    forms.sort(key=lambda form: form.get("updated_at") or "", reverse=True)
    return JSONResponse({"forms": forms})


@app.get("/forms/{form_id}")
async def get_form(form_id: str):
    """Read one saved form record by ID."""
    path = _form_file_path(form_id)
    if not os.path.exists(path):
        return JSONResponse({"error": f"Form '{form_id}' not found."}, status_code=404)

    try:
        with open(path, "r", encoding="utf-8") as file_obj:
            data = json.load(file_obj)
    except Exception as exc:
        return JSONResponse({"error": f"Failed to read form: {exc}"}, status_code=500)

    return JSONResponse(data)


@app.post("/forms/{form_id}/restore")
async def restore_form(form_id: str):
    """Load a saved form and restore it into shared agent state."""
    path = _form_file_path(form_id)
    if not os.path.exists(path):
        return JSONResponse({"error": f"Form '{form_id}' not found."}, status_code=404)

    try:
        with open(path, "r", encoding="utf-8") as file_obj:
            record = json.load(file_obj)
    except Exception as exc:
        return JSONResponse({"error": f"Failed to read form: {exc}"}, status_code=500)

    payload = {
        "peril": record.get("peril", {}),
        "questions": record.get("questions", []),
        "overall_outcome": record.get("overall_outcome", ""),
        "outcome_justification": record.get("outcome_justification", ""),
        "additional_analysis": record.get("additional_analysis"),
        "follow_ups": record.get("follow_ups"),
    }
    validation_error = _validate_form_payload(payload)
    if validation_error:
        return JSONResponse(
            {"error": f"Saved form is invalid: {validation_error}"}, status_code=500
        )

    state = APP_DEPS.state
    state.current_form_id = record.get("id", form_id)
    state.audit_form_result = payload
    state.audit_questions = payload["questions"]
    state.status = "complete"
    state.current_step = f"Restored saved form {state.current_form_id}"
    _upsert_audit_form_component(state, payload)

    return JSONResponse(
        {
            "message": "Form restored to state.",
            "form_id": state.current_form_id,
            "audit_form_result": state.audit_form_result,
        }
    )


@app.delete("/forms/{form_id}")
async def delete_form(form_id: str):
    """Delete a saved form from local JSON storage.

    If the deleted form is the currently active form, the reference is cleared
    from shared state.

    Args:
        form_id: Form identifier from the URL path.

    Returns:
        JSON confirmation or 404 if the form does not exist.
    """
    path = _form_file_path(form_id)
    if not os.path.exists(path):
        return JSONResponse({"error": f"Form '{form_id}' not found."}, status_code=404)

    try:
        os.remove(path)
    except Exception as exc:
        return JSONResponse({"error": f"Failed to delete form: {exc}"}, status_code=500)

    # Clear the active reference if this was the current form
    state = APP_DEPS.state
    if state.current_form_id == form_id:
        state.current_form_id = None

    return JSONResponse({"message": "Form deleted.", "form_id": form_id})


# =========================================================================
# Document summarization endpoint (NDJSON streaming)
# =========================================================================


class SummarizeDocPayload(BaseModel):
    """A single document payload sent by the frontend for summarization.

    Attributes:
        file_name: Original file name.
        content: Extracted text content of the document.
        mime_type: MIME type or extension string.
        document_type: High-level type classification (e.g. "Policy").
    """

    file_name: str
    content: str
    mime_type: str = "unknown"
    document_type: str = ""


class SummarizeRequest(BaseModel):
    """Request body for ``POST /summarize``.

    The frontend sends the document payloads directly so the endpoint
    is fully self-contained and does not read from shared AG-UI state.

    Attributes:
        documents: List of document payloads to summarize.
    """

    documents: list[SummarizeDocPayload]


@app.post("/summarize")
async def summarize_endpoint(body: SummarizeRequest):
    """Summarize documents one at a time, streaming NDJSON results.

    Processes each document sequentially — the LLM call for one document
    completes and its result is flushed to the client before the next
    document starts.  Each line is a self-contained JSON object (NDJSON).

    Args:
        body: Request with ``documents`` list.

    Returns:
        ``StreamingResponse`` with ``application/x-ndjson`` content type.
    """

    async def _generate():
        for doc in body.documents:
            if not doc.content.strip():
                line = json.dumps(
                    {
                        "file_name": doc.file_name,
                        "error": "No extractable content.",
                    }
                )
                yield line + "\n"
                continue

            try:
                prompt = format_document_summary_prompt(
                    file_name=doc.file_name,
                    document_content=doc.content,
                    file_type=doc.mime_type,
                    document_type=doc.document_type,
                )
                result = await document_summary_agent.run(prompt)
                payload = result.output.model_dump()
                payload["file_name"] = doc.file_name
                yield json.dumps(payload) + "\n"
            except Exception as exc:
                print(f"[SUMMARIZE ERROR] {doc.file_name}: {exc}", flush=True)
                line = json.dumps({"file_name": doc.file_name, "error": str(exc)})
                yield line + "\n"

    return StreamingResponse(
        _generate(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# =========================================================================
# Document search & sort endpoint
# =========================================================================


class SearchSortDocPayload(BaseModel):
    """A document payload with full metadata for the search/sort agent.

    Includes everything needed to construct a ``Document`` model so the
    agent can inspect metadata and full text via its tools.

    Attributes:
        file_name: Original file name.
        content_id: Unique content identifier.
        content: Extracted text content.
        mime_type: MIME type.
        content_url: URL or path for the document.
        claim_number: Associated claim number.
        domain: 'claim' or 'policy'.
        document_type: High-level type classification.
        document_sub_type: Finer-grained type.
        document_description: Human-readable description.
        create_date: Creation timestamp ISO string.
        source_system: Originating system.
        company_name: Associated company.
    """

    file_name: str
    content_id: str
    content: str = ""
    mime_type: str = "unknown"
    content_url: str = ""
    claim_number: str = ""
    domain: str = "claim"
    document_type: str = ""
    document_sub_type: str = ""
    document_description: str = ""
    create_date: str = ""
    source_system: str = ""
    company_name: str = ""


class SearchSortRequest(BaseModel):
    """Request body for ``POST /search-sort``.

    Attributes:
        query: The user's search or sort query.
        documents: Full document payloads including metadata and content.
    """

    query: str
    documents: list[SearchSortDocPayload]


@app.post("/search-sort")
async def search_sort_endpoint(body: SearchSortRequest):
    """Score documents against a user query using the search/sort agent.

    The agent reviews document metadata, optionally inspects top candidates,
    then returns a float score (0-1) and label for every document.

    Args:
        body: Request with ``query`` and ``documents`` list.

    Returns:
        JSON with ``scores`` list and a ``content_id_to_file_name`` mapping
        so the frontend can key results by file name.
    """
    if not body.documents:
        return JSONResponse({"scores": [], "content_id_to_file_name": {}})

    # Build Documents model from payloads for the agent's tools
    doc_models = []
    content_id_to_file_name: dict[str, str] = {}
    for d in body.documents:
        content_id_to_file_name[d.content_id] = d.file_name
        doc_models.append(
            Document(
                claimNumber=d.claim_number,
                contentId=d.content_id,
                mimeType=d.mime_type,
                contentURL=d.content_url or d.file_name,
                domain=d.domain if d.domain in ("claim", "policy") else "claim",
                documentType=d.document_type or None,
                documentSubType=d.document_sub_type or None,
                documentDescription=d.document_description or None,
                createDate=d.create_date or datetime.now(timezone.utc).isoformat(),
                sourceSystem=d.source_system or None,
                companyName=d.company_name or None,
                text=d.content,
            )
        )

    documents = Documents(documents=doc_models)
    deps = SearchSortDeps(documents=documents)

    try:
        prompt = format_doc_search_sort_prompt(query=body.query)
        result = await search_sort_agent.run(prompt, deps=deps)
        scores = result.output

        # Return scores + a mapping so the frontend can resolve content_id → file_name
        return JSONResponse(
            {
                "scores": [s.model_dump() for s in scores.scores],
                "content_id_to_file_name": content_id_to_file_name,
            }
        )

    except Exception as exc:
        print(f"[SEARCH-SORT ERROR] {exc}", flush=True)
        import traceback

        traceback.print_exc()
        return JSONResponse(
            {"error": str(exc), "scores": [], "content_id_to_file_name": {}},
            status_code=500,
        )


# =========================================================================
# Document tagging endpoint (batched NDJSON streaming)
# =========================================================================

MAX_BATCH_DOCS = 10
MAX_BATCH_CHARS = 25_000


def _batch_documents(
    docs: list[SummarizeDocPayload],
) -> list[list[SummarizeDocPayload]]:
    """Split documents into batches respecting doc-count and char-size limits.

    Args:
        docs: Full list of document payloads to partition.

    Returns:
        List of batches, each a list of ``SummarizeDocPayload``.
    """
    batches: list[list[SummarizeDocPayload]] = []
    current: list[SummarizeDocPayload] = []
    chars = 0

    for doc in docs:
        doc_chars = len(doc.content)
        if current and (len(current) >= MAX_BATCH_DOCS or chars + doc_chars > MAX_BATCH_CHARS):
            batches.append(current)
            current, chars = [], 0
        current.append(doc)
        chars += doc_chars

    if current:
        batches.append(current)
    return batches


class TagRequest(BaseModel):
    """Request body for ``POST /document-tags``.

    Attributes:
        documents: List of document payloads to tag.
    """

    documents: list[SummarizeDocPayload]


@app.post("/document-tags")
async def document_tags_endpoint(body: TagRequest):
    """Tag documents in batches, streaming NDJSON progress.

    Each document receives 1-4 tags from the predefined ``DocTag`` vocabulary.
    Pydantic validates tag values against the Literal type automatically, so
    no normalization pass is needed.

    Args:
        body: Request with ``documents`` list.

    Returns:
        ``StreamingResponse`` with ``application/x-ndjson`` content type.
    """

    async def _generate():
        batches = _batch_documents(body.documents)
        total_batches = len(batches)

        for batch_idx, batch in enumerate(batches):
            try:
                doc_dicts = [
                    {
                        "file_name": d.file_name,
                        "content": d.content,
                        "document_type": d.document_type,
                    }
                    for d in batch
                ]
                prompt = format_batch_tagger_prompt(documents=doc_dicts)
                result = await batch_tagger_agent.run(prompt)
                batch_results = result.output

                yield (
                    json.dumps(
                        {
                            "batch": batch_idx + 1,
                            "total_batches": total_batches,
                            "results": [r.model_dump() for r in batch_results.results],
                        }
                    )
                    + "\n"
                )

            except Exception as exc:
                print(f"[TAG BATCH ERROR] Batch {batch_idx + 1}: {exc}", flush=True)
                yield (
                    json.dumps(
                        {
                            "batch": batch_idx + 1,
                            "total_batches": total_batches,
                            "error": str(exc),
                        }
                    )
                    + "\n"
                )

        # Emit the static canonical tag list so the frontend can populate filter dropdowns
        yield (
            json.dumps(
                {
                    "done": True,
                    "canonical_tags": ALL_DOC_TAGS,
                }
            )
            + "\n"
        )

    return StreamingResponse(
        _generate(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# =========================================================================
# Text extraction helpers
# =========================================================================


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text and page count from a PDF using PyMuPDF.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Tuple of (extracted_text, page_count).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages), len(pages)


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text and approximate page count from a DOCX file.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Tuple of (extracted_text, estimated_page_count).
    """
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    # Rough page estimate: ~3000 chars per page
    page_estimate = max(1, len(text) // 3000)
    return text, page_estimate


def extract_text_from_xlsx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from an XLSX file by reading all sheets.

    Each row is joined with tabs, each sheet separated by a header line.

    Args:
        file_bytes: Raw XLSX bytes.

    Returns:
        Tuple of (extracted_text, sheet_count).
    """
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts), len(wb.sheetnames)


EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".xlsx": extract_text_from_xlsx,
}

EXT_TO_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


# =========================================================================
# Example / seed documents endpoint
# =========================================================================


@app.get("/example-docs")
async def example_docs_endpoint():
    """List pre-loaded example documents from the uploads directory.

    Scans ``UPLOAD_DIR`` for files with a supported extension, extracts text
    content using the same extractors as the upload endpoint, and returns a
    JSON array of document metadata + content.

    Returns:
        ``{"documents": [...]}`` with one entry per supported file.
    """
    docs: list[dict[str, Any]] = []
    try:
        for fname in sorted(os.listdir(UPLOAD_DIR)):
            ext = os.path.splitext(fname)[1].lower()
            if ext not in EXTRACTORS:
                continue
            fpath = os.path.join(UPLOAD_DIR, fname)
            if not os.path.isfile(fpath):
                continue

            file_bytes = open(fpath, "rb").read()  # noqa: SIM115
            text, pages = EXTRACTORS[ext](file_bytes)
            size = len(file_bytes)
            size_str = (
                f"{size / 1024 / 1024:.1f} MB" if size > 1024 * 1024 else f"{size / 1024:.1f} KB"
            )
            docs.append(
                {
                    "file_name": fname,
                    "mime_type": EXT_TO_MIME.get(ext, "application/octet-stream"),
                    "content": text,
                    "page_count": pages,
                    "file_size": size_str,
                    "file_size_bytes": size,
                    "path": f"/uploads/{fname}",
                }
            )
    except Exception as exc:
        print(f"[EXAMPLE-DOCS] Error scanning uploads dir: {exc}", flush=True)

    return JSONResponse({"documents": docs})


# =========================================================================
# File upload endpoint
# =========================================================================


@app.post("/upload")
async def upload_endpoint(file: UploadFile = File(...)):
    """Handle document file uploads with text extraction.

    Accepts a multipart ``file`` field. Supports .pdf, .docx, and .xlsx file
    types. Extracts text content so the agent can analyze the document.

    Args:
        file: The uploaded file provided via multipart form data.

    Returns:
        JSON with file metadata and extracted text content.
    """
    try:
        allowed_extensions = {".pdf", ".docx", ".xlsx"}
        filename = file.filename or "unknown"
        ext = os.path.splitext(filename)[1].lower()

        if ext not in allowed_extensions:
            return JSONResponse(
                {
                    "error": f"Unsupported file type: {ext}. Allowed: {', '.join(allowed_extensions)}"
                },
                status_code=400,
            )

        file_bytes = await file.read()

        # Save to disk (useful for debugging / future re-processing)
        file_path = os.path.join(UPLOAD_DIR, filename)
        with open(file_path, "wb") as f:
            f.write(file_bytes)

        # Extract text content
        extractor = EXTRACTORS.get(ext)
        if extractor:
            extracted_text, page_count = extractor(file_bytes)
        else:
            extracted_text, page_count = "", 0

        print(
            f"[UPLOAD] {filename}: {len(file_bytes)} bytes, "
            f"{page_count} pages, {len(extracted_text)} chars extracted",
            flush=True,
        )

        # Human-readable file size
        file_size = len(file_bytes)
        size_str = (
            f"{file_size / 1024 / 1024:.1f} MB"
            if file_size > 1024 * 1024
            else f"{file_size / 1024:.1f} KB"
        )

        return JSONResponse(
            {
                "filename": filename,
                "file_type": ext.lstrip("."),
                "file_size": size_str,
                "file_size_bytes": file_size,
                "page_count": page_count,
                "content": extracted_text,
                "path": file_path,
            }
        )
    except Exception as e:
        print(f"[UPLOAD ERROR] {e}", flush=True)
        import traceback

        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)


# =========================================================================
# Info / health endpoints
# =========================================================================


@app.get("/info")
async def info_endpoint():
    """Return agent information."""
    return JSONResponse(
        {
            "name": "Audit Assistant Agent",
            "version": "1.0.0",
            "protocol": "ag-ui",
            "description": "Analyze documents and generate custom audit questionnaires",
        }
    )


@app.get("/health")
async def health_endpoint():
    """Health check endpoint."""
    return JSONResponse(
        {
            "status": "healthy",
            "agent_ready": bool(os.getenv("OPENAI_API_KEY")),
        }
    )


# =========================================================================
# Doc Lens — text-to-image retrieval endpoints
# =========================================================================

# Lazy singleton: the CLIP embedder warmup is expensive (~5 s) so we only
# initialize on first actual use rather than at import / startup time.
_doc_lens_service = None


def _get_doc_lens_service():
    """Return (or create) the shared DocLensService singleton.

    Initialization is deferred to the first call so that the main FastAPI
    app starts quickly and the CLIP model is only loaded when needed.

    Returns:
        Fully initialised DocLensService instance.
    """
    global _doc_lens_service
    if _doc_lens_service is not None:
        return _doc_lens_service

    from doc_lens import (
        DocLensService,
        DuckDBStore,
        FastEmbedCLIPEmbedder,
        PDFExtractor,
        Settings,
    )

    settings = Settings()
    settings.ensure_dirs()

    db = DuckDBStore(settings.duckdb_path, embedding_dim=settings.embedding_dim)
    extractor = PDFExtractor(
        render_dpi=settings.render_dpi,
        min_area_ratio=settings.min_area_ratio,
        max_area_ratio=settings.max_area_ratio,
        crop_padding_px=settings.crop_padding_px,
    )
    embedder = FastEmbedCLIPEmbedder(
        model_key=settings.model_key,
        text_model_name=settings.text_model_name,
        image_model_name=settings.image_model_name,
        cache_dir=str(settings.fastembed_cache_dir),
    )

    _doc_lens_service = DocLensService(
        settings=settings,
        db=db,
        extractor=extractor,
        embedder=embedder,
    )
    return _doc_lens_service


# Mime types that Doc Lens can handle
DOC_LENS_PDF_MIMES = {"application/pdf"}
DOC_LENS_IMAGE_MIMES = {"image/jpeg", "image/png"}
DOC_LENS_ELIGIBLE_MIMES = DOC_LENS_PDF_MIMES | DOC_LENS_IMAGE_MIMES


class DocLensFilePayload(BaseModel):
    """A single file descriptor sent by the frontend for Doc Lens ingestion.

    Attributes:
        file_name: Original file name (must exist in the uploads directory).
        mime_type: MIME type string used to decide PDF vs image path.
    """

    file_name: str
    mime_type: str


class DocLensSessionRequest(BaseModel):
    """Request body for ``POST /doc-lens/session``.

    Attributes:
        files: List of file descriptors to ingest into the session.
    """

    files: list[DocLensFilePayload]


class DocLensQueryRequest(BaseModel):
    """Request body for ``POST /doc-lens/query``.

    Attributes:
        session_id: Active session identifier.
        query: Natural language query string.
        top_k: Maximum number of results (1-100).
        asset_types: Optional filter on asset types.
        document_ids: Optional filter on specific documents.
    """

    session_id: str
    query: str
    top_k: int = 10
    asset_types: list[str] | None = None
    document_ids: list[str] | None = None


class DocLensDocumentAssetsRequest(BaseModel):
    """Request body for ``POST /doc-lens/document-assets``.

    Attributes:
        session_id: Active session identifier.
        document_id: Document identifier to browse.
    """

    session_id: str
    document_id: str


@app.post("/doc-lens/session")
async def doc_lens_session_endpoint(body: DocLensSessionRequest):
    """Create a Doc Lens session and ingest files, streaming NDJSON progress.

    Each NDJSON line reports one of:
    - ``session_created`` with the new ``session_id``
    - ``ingest_start`` when a file begins processing
    - ``ingest_complete`` with per-file stats (page_count, assets, embeddings)
    - ``ingest_error`` if a single file fails (non-fatal, continues)
    - ``session_ready`` with the final session summary

    Args:
        body: Request with a list of file descriptors.

    Returns:
        StreamingResponse with ``application/x-ndjson`` content type.
    """
    session_id = str(uuid4())

    async def _generate():
        yield json.dumps({"type": "session_created", "session_id": session_id}) + "\n"

        svc = _get_doc_lens_service()
        total = len(body.files)

        for idx, file_desc in enumerate(body.files):
            file_index = idx + 1
            yield json.dumps({
                "type": "ingest_start",
                "file_name": file_desc.file_name,
                "mime_type": file_desc.mime_type,
                "file_index": file_index,
                "total_files": total,
            }) + "\n"

            try:
                file_path = os.path.join(UPLOAD_DIR, file_desc.file_name)
                if not os.path.isfile(file_path):
                    yield json.dumps({
                        "type": "ingest_error",
                        "file_name": file_desc.file_name,
                        "error": f"File not found: {file_desc.file_name}",
                        "file_index": file_index,
                        "total_files": total,
                    }) + "\n"
                    continue

                if file_desc.mime_type in DOC_LENS_PDF_MIMES:
                    result = svc.ingest_pdf(
                        session_id=session_id,
                        document_name=file_desc.file_name,
                        pdf_path=file_path,
                    )
                elif file_desc.mime_type in DOC_LENS_IMAGE_MIMES:
                    result = svc.ingest_image(
                        session_id=session_id,
                        document_name=file_desc.file_name,
                        image_path=file_path,
                    )
                else:
                    yield json.dumps({
                        "type": "ingest_error",
                        "file_name": file_desc.file_name,
                        "error": f"Unsupported mime type: {file_desc.mime_type}",
                        "file_index": file_index,
                        "total_files": total,
                    }) + "\n"
                    continue

                yield json.dumps({
                    "type": "ingest_complete",
                    "file_name": file_desc.file_name,
                    "file_index": file_index,
                    "total_files": total,
                    **result.model_dump(),
                }) + "\n"

            except Exception as exc:
                print(f"[DOC-LENS INGEST ERROR] {file_desc.file_name}: {exc}", flush=True)
                yield json.dumps({
                    "type": "ingest_error",
                    "file_name": file_desc.file_name,
                    "error": str(exc),
                    "file_index": file_index,
                    "total_files": total,
                }) + "\n"

        # Final summary
        try:
            summary = svc.get_session_summary(session_id)
            yield json.dumps({
                "type": "session_ready",
                **summary.model_dump(),
            }) + "\n"
        except Exception as exc:
            yield json.dumps({
                "type": "session_error",
                "error": f"Failed to get session summary: {exc}",
            }) + "\n"

    return StreamingResponse(
        _generate(),
        media_type="application/x-ndjson",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/doc-lens/query")
async def doc_lens_query_endpoint(body: DocLensQueryRequest):
    """Run a natural-language image query against an active Doc Lens session.

    Args:
        body: Query request with session_id, query text, and optional filters.

    Returns:
        JSON QueryResponse with ranked hits.
    """
    try:
        svc = _get_doc_lens_service()
        response = svc.query(
            session_id=body.session_id,
            query_text=body.query,
            top_k=body.top_k,
            asset_types=body.asset_types,
            document_ids=body.document_ids,
        )
        return JSONResponse(response.model_dump())
    except Exception as exc:
        print(f"[DOC-LENS QUERY ERROR] {exc}", flush=True)
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.post("/doc-lens/document-assets")
async def doc_lens_document_assets_endpoint(body: DocLensDocumentAssetsRequest):
    """List all extracted assets for one document in an active session.

    Args:
        body: Request with session_id and document_id.

    Returns:
        JSON object containing `session_id`, `document_id`, and `hits`.
    """
    try:
        svc = _get_doc_lens_service()
        hits = svc.list_document_assets(
            session_id=body.session_id,
            document_id=body.document_id,
        )
        return JSONResponse(
            {
                "session_id": body.session_id,
                "document_id": body.document_id,
                "hits": [hit.model_dump() for hit in hits],
            }
        )
    except Exception as exc:
        print(f"[DOC-LENS DOCUMENT-ASSETS ERROR] {exc}", flush=True)
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.get("/doc-lens/session/{session_id}")
async def doc_lens_session_summary(session_id: str):
    """Return the summary stats for a Doc Lens session.

    Args:
        session_id: Session identifier.

    Returns:
        JSON SessionSummary.
    """
    try:
        svc = _get_doc_lens_service()
        summary = svc.get_session_summary(session_id)
        return JSONResponse(summary.model_dump())
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.delete("/doc-lens/session/{session_id}")
async def doc_lens_clear_session(session_id: str):
    """Clear all data for a Doc Lens session.

    Args:
        session_id: Session identifier to purge.

    Returns:
        JSON confirmation.
    """
    try:
        svc = _get_doc_lens_service()
        svc.clear_session(session_id)
        return JSONResponse({"message": "Session cleared.", "session_id": session_id})
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


# =========================================================================
# Static file serving for uploaded documents (PDF viewer, downloads)
# =========================================================================

app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

# Doc Lens extracted assets — served so the frontend can <img src=...> them.
# Uses the asset_root from Settings which defaults to .cache/doc_lens_cache/assets.
_DOC_LENS_ASSET_DIR = os.path.join(
    os.path.dirname(__file__),
    ".cache", "doc_lens_cache", "assets",
)
os.makedirs(_DOC_LENS_ASSET_DIR, exist_ok=True)
app.mount(
    "/doc-lens-assets",
    StaticFiles(directory=_DOC_LENS_ASSET_DIR),
    name="doc-lens-assets",
)


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=BACKEND_PORT, reload=True)
